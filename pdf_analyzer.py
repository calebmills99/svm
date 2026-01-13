#!/usr/bin/env python3
"""
PDF Analyzer for Facebook Account Breach Evidence
Analyzes PDF files for security breach indicators and generates reports.

Author: Caleb Stewart
"""

import sys
import os
import json
import subprocess
from datetime import datetime
from typing import Dict, List, Any, Tuple


def get_recent_file_additions(num_files: int = 5) -> List[Tuple[str, str]]:
    """
    Get the N most recent file additions to the repository using Git.
    
    Args:
        num_files: Number of recent file additions to retrieve
        
    Returns:
        List of tuples (filepath, commit_date) for the most recent additions
    """
    try:
        # Get list of all added files with commit dates
        cmd = [
            'git', 'log', '--all', '--pretty=format:%ad|%H', '--date=iso',
            '--name-only', '--diff-filter=A'
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        
        # Parse output to get files with their addition dates
        files_with_dates = []
        lines = result.stdout.strip().split('\n')
        current_date = None
        
        for line in lines:
            if '|' in line:
                # This is a date|hash line
                current_date = line.split('|')[0]
            elif line.strip() and current_date:
                # This is a filename
                filepath = line.strip()
                # Only include files that still exist and are analyzable
                if os.path.exists(filepath):
                    ext = os.path.splitext(filepath)[1].lower()
                    if ext in ['.pdf', '.docx', '.doc', '.txt']:
                        files_with_dates.append((filepath, current_date))
        
        # Sort by date (most recent first) and return top N unique files
        files_with_dates.sort(key=lambda x: x[1], reverse=True)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_files = []
        for filepath, date in files_with_dates:
            if filepath not in seen:
                seen.add(filepath)
                unique_files.append((filepath, date))
                if len(unique_files) >= num_files:
                    break
        
        return unique_files
        
    except subprocess.CalledProcessError as e:
        print(f"Warning: Could not retrieve Git history: {e}")
        return []
    except Exception as e:
        print(f"Warning: Error getting recent additions: {e}")
        return []


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        import docx
        
        doc = docx.Document(docx_path)
        text_content = []
        
        print(f"Processing {len(doc.paragraphs)} paragraphs...")
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return "\n".join(text_content)
    
    except ImportError:
        print("\nError: python-docx library not available.")
        print("Please install it with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
        
        text_content = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            print(f"Processing {num_pages} pages...")
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                text_content.append(f"\n--- Page {page_num + 1} ---\n{text}")
        
        return "\n".join(text_content)
    
    except ImportError:
        print("PyPDF2 not installed. Attempting with pdfplumber...")
        try:
            import pdfplumber
            
            text_content = []
            with pdfplumber.open(pdf_path) as pdf:
                num_pages = len(pdf.pages)
                print(f"Processing {num_pages} pages...")
                
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"\n--- Page {page_num + 1} ---\n{text}")
            
            return "\n".join(text_content)
        
        except ImportError:
            print("\nError: No PDF library available.")
            print("Please install one of the following:")
            print("  pip install PyPDF2")
            print("  pip install pdfplumber")
            sys.exit(1)


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text content from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        print(f"Warning: Legacy DOC format not fully supported. Please convert to DOCX for best results.")
        # Try to process as DOCX anyway (some tools save as .doc but use DOCX format)
        try:
            return extract_text_from_docx(file_path)
        except Exception as e:
            print(f"Error: Could not process DOC file: {e}")
            return ""
    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ""
    else:
        print(f"Unsupported file type: {ext}")
        return ""


def analyze_content(text: str, user_name: str = "Caleb Stewart") -> Dict[str, Any]:
    """
    Analyze the extracted text for security indicators and evidence.
    
    Args:
        text: Extracted text content
        user_name: Name of the account owner
        
    Returns:
        Analysis results dictionary
    """
    analysis = {
        "user_name": user_name,
        "analysis_date": datetime.now().isoformat(),
        "security_indicators": [],
        "suspicious_activities": [],
        "ip_addresses": [],
        "login_locations": [],
        "timestamps": [],
        "key_findings": [],
        "recommendations": []
    }
    
    # Convert to lowercase for case-insensitive matching
    text_lower = text.lower()
    
    # Security indicators to look for
    security_keywords = [
        "unauthorized", "breach", "compromise", "hacked", "suspicious",
        "unrecognized", "unusual", "alert", "warning", "security",
        "failed login", "password reset", "account access", "ip address",
        "location", "device", "session", "login attempt"
    ]
    
    # Track found indicators
    for keyword in security_keywords:
        if keyword in text_lower:
            # Find context around the keyword
            idx = text_lower.find(keyword)
            context_start = max(0, idx - 100)
            context_end = min(len(text), idx + 100)
            context = text[context_start:context_end].strip()
            
            analysis["security_indicators"].append({
                "keyword": keyword,
                "context": context
            })
    
    # Look for IP addresses
    import re
    ip_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'
    potential_ips = re.findall(ip_pattern, text)
    
    # Validate that each octet is in valid range (0-255)
    valid_ips = []
    for ip in potential_ips:
        octets = ip.split('.')
        if all(0 <= int(octet) <= 255 for octet in octets):
            valid_ips.append(ip)
    
    analysis["ip_addresses"] = list(set(valid_ips))  # Remove duplicates
    
    # Look for geographic locations
    location_keywords = ["location:", "from:", "country:", "city:", "region:"]
    for line in text.split('\n'):
        for loc_key in location_keywords:
            if loc_key in line.lower():
                analysis["login_locations"].append(line.strip())
    
    # Look for timestamps using dateutil for better parsing
    try:
        from dateutil import parser as date_parser
        
        # Find potential date strings
        date_patterns = [
            r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b',  # MM/DD/YYYY or DD/MM/YYYY
            r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b',     # YYYY/MM/DD
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b',  # Month DD, YYYY
        ]
        
        potential_dates = []
        for pattern in date_patterns:
            potential_dates.extend(re.findall(pattern, text, re.IGNORECASE))
        
        # Validate dates using dateutil
        valid_dates = []
        for date_str in potential_dates:
            try:
                parsed_date = date_parser.parse(date_str, fuzzy=False)
                valid_dates.append(date_str)
            except (ValueError, date_parser.ParserError):
                # Skip invalid dates
                pass
        
        analysis["timestamps"] = list(set(valid_dates[:20]))  # Limit to first 20 unique valid dates
    
    except ImportError:
        # Fallback to basic regex if dateutil not available
        date_pattern = r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b|\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b'
        dates = re.findall(date_pattern, text)
        analysis["timestamps"] = list(set(dates[:20]))
    
    # Generate key findings
    if analysis["security_indicators"]:
        analysis["key_findings"].append(
            f"Found {len(analysis['security_indicators'])} security-related indicators in the document"
        )
    
    if analysis["ip_addresses"]:
        analysis["key_findings"].append(
            f"Identified {len(analysis['ip_addresses'])} unique IP addresses"
        )
    
    if analysis["login_locations"]:
        analysis["key_findings"].append(
            f"Found {len(analysis['login_locations'])} location references"
        )
    
    # Generate recommendations
    analysis["recommendations"] = [
        "Review all IP addresses and verify if they match your known locations",
        "Check timestamps for any activity during times when you were not active",
        "Look for any unrecognized devices or login locations",
        "Document all suspicious activities for potential legal/security reporting",
        "Consider enabling two-factor authentication if not already enabled",
        "Review and update security settings on all associated accounts"
    ]
    
    return analysis


def generate_report(analysis: Dict[str, Any], output_path: str = "analysis_report.txt"):
    """
    Generate a human-readable report from the analysis.
    
    Args:
        analysis: Analysis results dictionary
        output_path: Path for the output report file
    """
    report_lines = [
        "=" * 80,
        "FACEBOOK ACCOUNT BREACH EVIDENCE ANALYSIS REPORT",
        "=" * 80,
        f"\nUser: {analysis['user_name']}",
        f"Analysis Date: {analysis['analysis_date']}",
        "\n" + "=" * 80,
    ]
    
    # Key Findings Section
    report_lines.append("\nKEY FINDINGS:")
    report_lines.append("-" * 80)
    if analysis["key_findings"]:
        for i, finding in enumerate(analysis["key_findings"], 1):
            report_lines.append(f"{i}. {finding}")
    else:
        report_lines.append("No significant findings detected.")
    
    # Security Indicators Section
    report_lines.append("\n\nSECURITY INDICATORS:")
    report_lines.append("-" * 80)
    if analysis["security_indicators"]:
        # Group by keyword to avoid repetition
        keyword_groups = {}
        for indicator in analysis["security_indicators"]:
            keyword = indicator["keyword"]
            if keyword not in keyword_groups:
                keyword_groups[keyword] = []
            keyword_groups[keyword].append(indicator["context"])
        
        for keyword, contexts in keyword_groups.items():
            report_lines.append(f"\n'{keyword.upper()}' (found {len(contexts)} time(s)):")
            # Show first 3 contexts for each keyword
            for context in contexts[:3]:
                report_lines.append(f"  Context: ...{context}...")
    else:
        report_lines.append("No security indicators found.")
    
    # IP Addresses Section
    report_lines.append("\n\nIP ADDRESSES DETECTED:")
    report_lines.append("-" * 80)
    if analysis["ip_addresses"]:
        for ip in analysis["ip_addresses"][:50]:  # Limit to 50 IPs
            report_lines.append(f"  - {ip}")
    else:
        report_lines.append("No IP addresses found.")
    
    # Locations Section
    report_lines.append("\n\nLOCATION REFERENCES:")
    report_lines.append("-" * 80)
    if analysis["login_locations"]:
        for location in analysis["login_locations"][:20]:  # Limit to 20 locations
            report_lines.append(f"  - {location}")
    else:
        report_lines.append("No location references found.")
    
    # Timestamps Section
    report_lines.append("\n\nTIMESTAMPS FOUND:")
    report_lines.append("-" * 80)
    if analysis["timestamps"]:
        for timestamp in analysis["timestamps"]:
            report_lines.append(f"  - {timestamp}")
    else:
        report_lines.append("No timestamps found.")
    
    # Recommendations Section
    report_lines.append("\n\nRECOMMENDATIONS:")
    report_lines.append("-" * 80)
    for i, rec in enumerate(analysis["recommendations"], 1):
        report_lines.append(f"{i}. {rec}")
    
    # Implications Section
    report_lines.append("\n\nIMPLICATIONS:")
    report_lines.append("-" * 80)
    implications = [
        "SECURITY BREACH EVIDENCE: This document may contain evidence of unauthorized",
        "access to your Facebook account. The presence of security indicators suggests",
        "that your account may have been compromised.",
        "",
        "LEGAL CONSIDERATIONS: This evidence may be relevant for:",
        "  - Filing a complaint with Meta/Facebook",
        "  - Reporting to law enforcement or cybersecurity authorities",
        "  - Identity theft protection services",
        "  - Legal proceedings if damages occurred",
        "",
        "IMMEDIATE ACTIONS: Consider taking the following steps:",
        "  - Change passwords on all accounts (especially if reused)",
        "  - Enable two-factor authentication",
        "  - Review account activity and settings",
        "  - Monitor credit reports for identity theft",
        "  - Report the breach to appropriate authorities",
        "",
        "DOCUMENTATION: Keep this analysis and the original PDF as evidence.",
        "Consider consulting with a cybersecurity professional or attorney for",
        "proper handling of this evidence."
    ]
    
    for line in implications:
        report_lines.append(line)
    
    report_lines.append("\n" + "=" * 80)
    report_lines.append("END OF REPORT")
    report_lines.append("=" * 80)
    
    # Write to file
    report_content = "\n".join(report_lines)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(report_content)
    
    print(f"\nReport generated: {output_path}")
    return report_content


def process_file(file_path: str, user_name: str = "Caleb Stewart") -> None:
    """
    Process a single file for breach evidence analysis.
    
    Args:
        file_path: Path to the file to analyze
        user_name: Name of the account owner
    """
    print(f"\n{'='*80}")
    print(f"Processing: {file_path}")
    print(f"{'='*80}")
    
    # Extract text from file
    print(f"\nExtracting text from: {file_path}")
    text_content = extract_text_from_file(file_path)
    
    if not text_content:
        print(f"Warning: No text could be extracted from {file_path}")
        return
    
    # Save extracted text
    base_name = os.path.splitext(file_path)[0]
    text_output = f"{base_name}_extracted_text.txt"
    with open(text_output, 'w', encoding='utf-8') as f:
        f.write(text_content)
    print(f"Extracted text saved to: {text_output}")
    
    # Analyze content
    print("\nAnalyzing content for security indicators...")
    analysis = analyze_content(text_content, user_name=user_name)
    
    # Save analysis as JSON
    json_output = f"{base_name}_analysis.json"
    with open(json_output, 'w', encoding='utf-8') as f:
        json.dump(analysis, f, indent=2)
    print(f"Analysis data saved to: {json_output}")
    
    # Generate report
    print("\nGenerating report...")
    report_output = f"{base_name}_report.txt"
    generate_report(analysis, report_output)
    
    # Print summary
    print("\n" + "=" * 80)
    print("ANALYSIS SUMMARY")
    print("=" * 80)
    print(f"File: {file_path}")
    print(f"Security Indicators Found: {len(analysis['security_indicators'])}")
    print(f"Unique IP Addresses: {len(analysis['ip_addresses'])}")
    print(f"Location References: {len(analysis['login_locations'])}")
    print(f"Timestamps Found: {len(analysis['timestamps'])}")
    print("=" * 80)


def main():
    """Main function to run the file analyzer."""
    print("=" * 80)
    print("BREACH EVIDENCE ANALYZER - SCANNING RECENT ADDITIONS")
    print("=" * 80)
    
    # Check if user wants to scan recent additions or a specific file
    if len(sys.argv) > 1 and sys.argv[1] == "--recent":
        # Scan the 5 most recent additions
        print("\nScanning the 5 most recent file additions to the repository...")
        recent_files = get_recent_file_additions(5)
        
        if not recent_files:
            print("\nNo recent file additions found or Git not available.")
            print("Falling back to scanning PDF/DOCX files in current directory...")
            # Fall back to scanning current directory
            files_to_scan = []
            for ext in ['.pdf', '.docx', '.doc']:
                files_to_scan.extend([f for f in os.listdir('.') if f.endswith(ext)])
            recent_files = [(f, "unknown") for f in files_to_scan[:5]]
        
        if not recent_files:
            print("\nNo files found to scan.")
            print("\nUsage:")
            print("  python pdf_analyzer.py --recent              # Scan 5 most recent additions")
            print("  python pdf_analyzer.py [filename]            # Scan specific file")
            print("  python pdf_analyzer.py                       # Scan first PDF in directory")
            sys.exit(1)
        
        print(f"\nFound {len(recent_files)} recent file(s) to analyze:")
        for i, (filepath, date) in enumerate(recent_files, 1):
            print(f"  {i}. {filepath} (added: {date})")
        
        # Get user name if provided
        user_name = sys.argv[2] if len(sys.argv) > 2 else "Caleb Stewart"
        
        # Process each file
        print(f"\n{'='*80}")
        print(f"Processing {len(recent_files)} file(s)...")
        print(f"{'='*80}")
        
        for filepath, date in recent_files:
            if os.path.exists(filepath):
                process_file(filepath, user_name)
            else:
                print(f"\nWarning: File not found: {filepath}")
        
        print(f"\n{'='*80}")
        print("SCAN COMPLETE")
        print(f"{'='*80}")
        print(f"\nAnalyzed {len(recent_files)} file(s).")
        print("Review the generated *_report.txt files for detailed findings.")
        print(f"{'='*80}")
        
    else:
        # Original behavior: process a single file
        # Check for PDF files in the current directory
        pdf_files = [f for f in os.listdir('.') if f.endswith('.pdf')]
        
        if not pdf_files:
            print("\nNo PDF files found in the current directory.")
            print("\nUsage:")
            print("  python pdf_analyzer.py --recent              # Scan 5 most recent additions")
            print("  python pdf_analyzer.py [filename.pdf]        # Scan specific file")
            print("  python pdf_analyzer.py                       # Scan first PDF in directory")
            print("\nExamples:")
            print("  python pdf_analyzer.py --recent")
            print("  python pdf_analyzer.py facebook_data.pdf")
            print("  python pdf_analyzer.py facebook_data.pdf \"Caleb Stewart\"")
            
            # Check if a filename was provided as argument
            if len(sys.argv) > 1:
                pdf_path = sys.argv[1]
                if not os.path.exists(pdf_path):
                    print(f"\nError: File '{pdf_path}' not found.")
                    sys.exit(1)
            else:
                sys.exit(1)
        else:
            # Use provided filename or first PDF found
            if len(sys.argv) > 1:
                pdf_path = sys.argv[1]
                if not os.path.exists(pdf_path):
                    print(f"\nError: File '{pdf_path}' not found.")
                    sys.exit(1)
            else:
                pdf_path = pdf_files[0]
                print(f"\nFound PDF files: {', '.join(pdf_files)}")
                print(f"Processing: {pdf_path}")
        
        # Allow user to specify name via command line argument, default to "Caleb Stewart"
        user_name = sys.argv[2] if len(sys.argv) > 2 else "Caleb Stewart"
        
        # Process the file
        process_file(pdf_path, user_name)


if __name__ == "__main__":
    main()
